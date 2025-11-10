# This Python script has been created to ingest the output of the "DICK.py" script
# to create a Word Document containing all relevant information that can be auto-filled
# for an "As-Built" Cribl document.
# Files were created by the TSI C4 cohort as of 21/12/2023.
 
 
# Import docx library and dependencies
# All information for the docx library was found at "https://python-docx.readthedocs.io/en/latest/"
# If run locally, the docx library needs to be installed with "pip install python-docx"
from docx import Document
from docx.shared import Inches
 
# Used for getting todays date
from datetime import datetime, timezone

from math import ceil
# Used for getting client and consultant team information
import csv,json
 
# Imports the API call script for getting information
import API_calls as get_json

import requests
 
# Function to populate the as built document
def populate_document():
    ## Page 1
    title_page()
 
    ## Page 2
    change_control_table()
    version_control_table()
   
    ## Page 3
    executive_summary(executive_summary_text)
   
    ## Page 4
    project_management()
    customer_team()
 
    ## Page 5+
    architecture()
    worker_groups(group_names)
    leader_node()
    leader_load_balance()
    worker_specs()
    #git_control()
    create_all_autofill_tables()
    use_cases()
    block_footer_creation()
   
# Function to populate health document
def health_document():
    health_title()
    health_document_control()

    health_executive_summary()
    health_current_configurations()
    health_reference()

# Function writing the title for the document
def health_title():
    document_obj.add_heading(client_name+" Stream Health Check",0)
    document_obj.add_heading("Cribl Services Documentation",1)
    document_obj.add_page_break()

# Function creating table and filling information for the document control and history
def health_document_control():
    document_control_information=(
        ("Description", "This document will cover the health chekc findings that were performed by the Services Engineer"),
        ("Intended Audience", "Sales, Services Consultants, Sales Engineers, TSE, CSM and End Customers"),
        ("Updated", "1/29/2024"),
        ("Document Status", "V1.0"),
        ("Author", consultant_name)
    )
    document_obj.add_heading("Document Control", 1)
    document_table=document_obj.add_table(rows=0, cols=2)
    table_by_record(document_table, document_control_information, "Colorful Grid Accent 1")
    document_history=(
        ("Version", "Date Issued", "Status", "Description"),
        ("1.0", datetime.now().strftime("%m-%d-%Y"), "Complete", "Automated draft")
    )
    document_obj.add_heading("Document History", 2)
    history_table=document_obj.add_table(rows=0,cols=4)
    table_by_record(history_table, document_history, "Colorful Grid Accent 1")
    document_obj.add_page_break()

# Function creating tables and filling information for leaders, workers, and their best practices
def health_executive_summary():
    document_obj.add_heading("Executive Summary", 0)
    document_obj.add_heading("Architecture", 1)
    document_obj.add_heading("Overview",2)

    # Git Summary Table
    git_authType=""
    git_action=""
    git_remote=""
    git_message=""
    for i, d in api_return.items():
        if i!="Git":
            continue
        for key in d:
            git_authType=key["AuthType"]
            git_action=key["AutoAction"]
            git_remote=key["Remote"]
            git_message=key["Commit Message"]
    git_info=(
        ("AuthType",git_authType),
        ("AutoAction",git_action),
        ("Remote",git_remote),
        ("Commit Message", git_message)
    )
    document_obj.add_heading("Git Authentication",1)
    git_table=document_obj.add_table(rows=0,cols=2)
    table_by_record(git_table, git_info, "Medium Grid 1 Accent 1")

    # Leader Summary
    active_workers=""
    cpu="" 
    memory=""
    disk=""
    for i,d in api_return.items():
        if i!="Leader Node":
            continue
        for key in d:
            active_workers=key["Active Workers/Total Workers"]
            cpu=key["CPU"]
            memory=key["Total Memory in GB"]
            disk=key["Disk Total in TB"]
            hostname=key["Hostname"]
            ip=key["IP Address"]
    health_leader_info=(
        ("Hostname", hostname),
        ("IP Address", ip),
        ("Workers Reporting", str(active_workers)),
        ("Non-Local Authentication", "Enter Here"),
        ("Memory", str(memory)+" GB's"),
        ("CPU", str(cpu)),
        ("Disk Usage",str(disk)+" Tb's"),
        ("Network Utilization", "Enter Here"),
        ("Log WARN/ERROR", "Enter Here")
    )
    document_obj.add_heading("Leader System Specification",1)
    leader_table=document_obj.add_table(rows=0,cols=2)
    table_by_record(leader_table, health_leader_info, "Medium Grid 1 Accent 1")

    # Leader Best Practices
    leader_best=(
        ("Boot-Start Enabled", "Enter Here"),
        ("Non-Root user", "Enter Here"),
        ("Git Repo Configured and Working", "Enter Here"),
        ("Default Worker Group Used and Logging Sent to Analytics", "Enter Here"),
        ("Platform", "Enter Here"),
        ("NTP Synced", "Enter Here"),
        ("ULimits Properly Set", "Enter Here")
    )
    document_obj.add_heading("Best Practice", 2)
    leader_best_table=document_obj.add_table(rows=0,cols=2)
    table_by_record(leader_best_table, leader_best, "Medium Grid 1 Accent 1")

    # Worker Summary
    total_workers=""
    total_cpu=""
    total_memory=""
    total_disk=""
    total_disk_usage=""
    for i,d in api_return.items():
        if i!="Workers":
            continue
        for key in d:
            total_workers=key["Worker Count"]
            total_cpu=key["Total Cores"]
            total_memory=key["Total Memory"]
            total_disk=key["Total Disk"]
            total_disk_usage=key["Disk Usage"]
    worker_summary=(
        ("Total Count Ratio", str(total_workers)),
        ("Total CPU/Memory Load", str(total_cpu)+"/"+str(total_memory)),
        ("File Descriptors", "Enter Here"),
        ("Total Disk Usage", str(total_disk_usage)+"/"+str(total_disk)+" GB's"),
        ("Network Utilization", "Enter Here")
    )
    document_obj.add_heading("Workers", 1)
    worker_sum_table=document_obj.add_table(rows=0,cols=2)
    table_by_record(worker_sum_table, worker_summary, "Medium Grid 1 Accent 1")

    # Worker Best Practices
    worker_best=(
        ("PQ Separate Volume", "Enter Here"),
        ("Boot-Start Enabled", "Enter Here"),
        ("Version Match", "Enter Here"),
        ("Non-Root User", "Enter Here"),
        ("NTP Synced", "Enter Here"),
        ("ULimits Properly Set", "Enter Here")
    )
    document_obj.add_heading("Best Practice", 2)
    worker_best_table=document_obj.add_table(rows=0,cols=2)
    table_by_record(worker_best_table, worker_best, "Medium Grid 1 Accent 1")

    # Data Flow
    document_obj.add_heading("Data Flow", 1)
    document_obj.add_paragraph("Leverage Cribl compression to reduce cloud egress costs")

# Function creating tables and filling information for the configurations of Sources, Collectors, Destinations, Pipelines, and Routes for each worker group
def health_current_configurations():
    # Current Configurations
    document_obj.add_heading("Current Configurations", 0)
    
    #Workers Config
    document_obj.add_heading("Workers", 1)
    document_obj.add_heading("API Logs", 2)
    document_obj.add_heading("Warnings and Errors", 3)
    channel=""
    msg=""
    workerID=""
    for i, d in api_return.items():
        if i!="API_Logs":
            continue
        try:
            for key in d:
                channel=key["Channel"]
                msg=key["Message"]
                try:
                    workerID=key["Worker_ID"]
                except KeyError:
                    workerID="N/A"
                document_obj.add_paragraph("Error for | Worker ID: "+workerID+" | Channel: "+channel+" | Message: "+msg)
        except:
            document_obj.add_paragraph("No errors found in this environment")

    document_obj.add_heading("Worker Process Logs", 2)
    document_obj.add_heading("Warnings and Errors", 3)
    document_obj.add_paragraph("Enter warnings and errors here")

    #Sources config
    document_obj.add_heading("Sources", 1)
    document_obj.add_heading("Connectivity",2)
    for i,d in api_return.items():
        if i!="Sources":
            continue
        for key in d:
            all_green=True
            document_obj.add_heading(key+" Worker Group", 3)
            for v in d[key]:
                if v["Health"]!="Green":
                    all_green=False
                    document_obj.add_paragraph(v["Source Name"]+" has failed, check for more information")
            if all_green==True:
                document_obj.add_paragraph("All Sources showing green")

    #Collectors config
    document_obj.add_heading("Collectors", 1)
    document_obj.add_heading("Job Failures", 2)
    for i,d in api_return.items():
        all_green=True
        if i!="Jobs":
            continue
        for key in d:
            exists=False
            document_obj.add_heading(key+" Worker Group", 3)
            for v in d[key]:
                exists=True
                if v["Status"]=="failed":
                    all_green=False
                    document_obj.add_paragraph(v["Job ID"]+" tests has failed, check for more information")
            if exists==True:
                if all_green==True:
                    document_obj.add_paragraph("All collectors jobs showing green")
            else:
                document_obj.add_paragraph("No collectors configured")

    #Destinations config
    document_obj.add_heading("Destinations", 1)
    document_obj.add_heading("Connectivity", 2)
    for i,d in api_return.items():
        if i!="Destinations":
            continue
        for key in d:
            all_green=True
            document_obj.add_heading(key+" Worker Group",3)
            for v in d[key]:
                if v["Destination Name"]=="default":
                    continue
                if v["Health"]=="Red":
                    all_green=False
                    document_obj.add_paragraph(v["Destination Name"]+" destination has failed, check for more information")
            if all_green==True:
                document_obj.add_paragraph("All Destinations showing green")
        document_obj.add_heading("Back Pressure/Queues", 2)
        for key in d:
            document_obj.add_heading(key+" Worker Group", 3)
            populate_table_by_json(i,key)

    #Pipelines config
    document_obj.add_heading("Pipelines", 1)
    for i, d in api_return.items():
        if i!= "Pipelines":
            continue
        for key in d:
            document_obj.add_heading(key+" Worker Group", 3)
            populate_table_by_json(i, key)

    #Routes config
    document_obj.add_heading("Routes", 1)
    for i, d in api_return.items():
        if i!= "Routes":
            continue
        for key in d:
            document_obj.add_heading(key+" Worker Group", 3)
            populate_table_by_json(i, key)

    #Packs config
    document_obj.add_heading("Packs", 1)
    for i, d in api_return.items():
        if i!="Packs":
            continue
        for key in d:
            document_obj.add_heading(key+" Worker Group", 3)
            populate_table_by_json(i, key)

# Prints out all the references for the health checks in bullet points
def health_reference():

    document_obj.add_heading("Reference Documents", 0)
    document_obj.add_paragraph("https://docs.cribl.io/docs ", style="List Bullet")
    document_obj.add_paragraph("Current LogStream version's docs", style="List Bullet 2")
    document_obj.add_paragraph("New docs only created with major and minor releases (X.Y.Z, where X is the major and Y is the minor release)", style="List Bullet 2")
    document_obj.add_paragraph("Docs for past releases are available using https://docs.cribl.io/vX.Y/docs URL structure", style="List Bullet 2")

    document_obj.add_paragraph("https://docs.cribl.io/docs/introduction-reference ", style="List Bullet")
    document_obj.add_paragraph("Contains 1st- and 3rd-party docs for JavaScript expressions (String, date, and number manipulation, among others", style="List Bullet 2")

    document_obj.add_paragraph("https://docs.cribl.io/docs/api-docs ", style="List Bullet")
    document_obj.add_paragraph("Contains Swagger-based documentation for the LogStream API", style="List Bullet 2")
    
    document_obj.add_paragraph("https://docs.cribl.io/ ", style="List Bullet")
    document_obj.add_paragraph("Links to PDF of full documentation for current LogStream release", style="List Bullet 2")

# Create Title page with heading, cribl logo and page break
def title_page():
    document_obj.add_heading(client_name + " - Cribl Stream As-Built Architecture", 0)
    document_obj.add_picture("Scripts/cribl_logo.png", width=Inches(2.5))
    document_obj.add_page_break()
 
# Create and populate control table
def change_control_table():
    # Tuple of information used for the change table
    change_table_information = (
        ("Title:", client_name + " - Cribl Stream As Built Architecture"),
        ("Version", "1.0"),
        ("Date of version", (datetime.now(timezone.utc)).strftime("%D")),
        ("Created by", consultant_name)
    )
    document_obj.add_heading("Change Control", 1)
    # Creates a table with an extra line for formatting
    change_table = document_obj.add_table(rows=0, cols=2)
    table_by_record(change_table, change_table_information, "Colorful Grid Accent 1")
 
# Create and populate version table
def version_control_table():
    # Tuple of information used for the version table
    version_table_information = (
        ("Date", "Version", "Created By", "Description of Change"),
        (datetime.now().strftime("%d-%m-%Y"), "1.0", consultant_name, "Initial Draft")
    )
    document_obj.add_heading("Version History", 2)
    version_table = document_obj.add_table(rows=0, cols=4)
    table_by_record(version_table, version_table_information, "Colorful Grid Accent 1")
 
# Populates table for number of lines
def table_by_record(table, record, style_name):
    # Ingests table obj, record tuple, style_name str
    for record_line in record:
        # Adds a row to the table for each record
        table_rows = table.add_row().cells
        # Populates the table with the information within the tuple
        for index,entry in enumerate(record_line):
            table_rows[index].text = entry
    # Applies the provided table style
    table.style = style_name

# Creates a table that displays all information about the workers in the environment
def worker_specs():
    document_obj.add_heading("Workers", 1)
    for group in get_json.all_workgroups(header,url):
        # Iterates through each worker
        if group in api_return["Workers"]:
            document_obj.add_heading(group + " Workers", 2)
            populate_table_by_json("Workers", group)
 
# Create a bullet point list based on a csv for roles
# Ingests a variable containing lines of a csv as a list
def bullet_point_by_csv(csv_file_lines):
    # Iterates for each row of the csv
    for index, row in enumerate(csv_file_lines):
        # Allow for auto replace of <consultant> in the file to be adapted if need be
        # Creates a bullet point list entry for each entry in the passed csv
        # Note: The docx library does not provide a more succinct way to add bullet points, I'm sorry :(
        if str(row["name"]) == "<CONSULTANT>":
            document_obj.add_paragraph(consultant_name + " - " + str(row["role"]), style="List Bullet")
        else:
            document_obj.add_paragraph(str(row["name"]) + " - " + str(row["role"]), style="List Bullet")
 
# Add executive summary heading and information
def executive_summary(summary_text):
    document_obj.add_heading("Executive Summary", 2)
    document_obj.add_paragraph(summary_text)
    document_obj.add_paragraph("\n<INSERT TABLE OF CONTENTS>")
    document_obj.add_page_break()
 
# Add project management section with bullet pointed list of cribl team
def project_management():
    # Opens the "cribl_team.csv" for the information of the team
    cribl_team_csv = open("Scripts/cribl_team.csv", newline="")
    cribl_team_csv_reader = csv.DictReader(cribl_team_csv, fieldnames = ["name","role"])
    next(cribl_team_csv_reader, None)
    # Adds headings and text for the project management section
    document_obj.add_heading("Project Management", 1)
    document_obj.add_heading("Cribl Team", 2)
    document_obj.add_paragraph("The Following key personnel make up the account team on the Cribl side to ensure the success of this deployment.")
    # Creates the bullet point list of names from the csv
    bullet_point_by_csv(cribl_team_csv_reader)
    cribl_team_csv.close()
 
# Populates Client Team information as a list
def customer_team():
    # Opens the "customer_team.csv" for the information of the team
    customer_team_csv = open("Scripts/customer_team.csv", newline="")
    customer_team_csv_reader = csv.DictReader(customer_team_csv, fieldnames = ["name","role"])
    next(customer_team_csv_reader, None)
    # Adds heading and text for the account team
    document_obj.add_heading(client_name + " Team", 2)
    document_obj.add_paragraph("The Following key personnel make up the account team on the " + client_name + " side to ensure the success of this deployment.")
    # Creates the bullet point list of names from the csv
    bullet_point_by_csv(customer_team_csv_reader)
    customer_team_csv.close()
    document_obj.add_page_break()
 
# Populates simple achitecture section information
def architecture():
    document_obj.add_heading("Architecture", 1)
    document_obj.add_heading("Deployment Overview", 2)
    document_obj.add_paragraph(client_name + " purchased Cribl Stream to <INSERT REASON FOR CRIBL>.")
    document_obj.add_heading("Data Centers", 3)
    document_obj.add_paragraph("\t-")

# Function that takes the values of the envrionments git settings and prints them out
def git_control():
    document_obj.add_heading("Git Integration", 1)
    group_table=document_obj.add_table(rows=0, cols=2)
    # This loop will take all information from the git settings and assign them to 4 values
    for i,d in api_return.items():
        if i!="Git":
            continue
        for key in d:
            authType=key["AuthType"]
            autoAction=key["AutoAction"]
            remote=key["Remote"]
            commitMessage=key["Commit Message"]
    git_information=(
        ("AuthType", authType),
        ("AutoAction", autoAction),
        ("Remote", remote),
        ("Commit Message",commitMessage)
    )
    table_by_record(group_table, git_information, "Medium Grid 1 Accent 1")

# A function to add the names of each worker group with default text for filling in the use of the 
# group and for creating tables that contain the number of Workers, Routes, Sources, Destinations
# and Pipelines for each group. 
# There is a limit for up to 5 groups per table to prevent overcrowding in the tables.
def worker_groups(list_of_groups):
    # Establishes the items to count for each group in the table
    row_names = ("Workers", "Routes", "Sources", "Destinations", "Pipelines")
    document_obj.add_heading("Worker Groups", 3)
    # Creates the group name header with <USE OF WORKER GROUP> underneath to be completed
    for group_name in list_of_groups:
        document_obj.add_paragraph("- " + group_name)
        document_obj.add_paragraph("\t- <USE OF WORKER GROUP>")
    # Creates the tables for the count information
    # Repeats the process for each time the group exceeds vvvv this number
    for group_block_num in range(ceil(len(list_of_groups) / 5)):
        # Identifies the start and end of the group list
        start_index = group_block_num*5
        end_index = start_index + 5
        # Catches the end index if it exceeds the length of groups
        if end_index > len(list_of_groups):
            end_index = len(list_of_groups)
        # Creates a table object with as many rows as needed for the different items to count add 1 
        # for headings and columns as the groups it is dealing with in this loop add 1 for Row names.
        group_table = document_obj.add_table(rows=len(row_names) + 1, cols=(end_index - start_index) + 1)
        title_line = group_table.rows[0].cells
        # Sets the column titles to the group names, offset by 1
        for index,group_name in enumerate(list_of_groups[start_index:end_index]):
            title_line[index+1].text = group_name
        # Iterates through each item to be counted and sets the row name
        for index,row_name in enumerate(row_names):
            working_line = group_table.rows[index+1].cells
            working_line[0].text="Number of "+row_name
            # Iterates through the index of up to 5 groups
            for index2,name in enumerate(list_of_groups[start_index:end_index]):
                # Try is used to catch when there is no data for a specified item
                try:
                    # Sets the cell of the table equal to a count of the item for the specified group
                    working_line[index2+1].text = str(len(api_return[row_name][name]))    
                except KeyError:
                    # No data exception handling sets the cell to "N/A"
                    working_line[index2+1].text = "N/A"
        group_table.style="Colorful Grid Accent 1"
        document_obj.add_paragraph()

# Creates a table for Leader Nodes information 
def leader_node():
    version=""
    os=""
    cpu=""
    memory=""
    for i,d in api_return.items():
        if i != "Leader Node":
            continue
        for key in d:
            version=key["Cribl Version"]
            os=key["OS"]
            cpu=key["CPU"]
            memory=key["Total Memory in GB"]
            file=key["Cribl Home"]
            ip=key["IP Address"]
            host=key["Hostname"]

    leader_node_information=(
        ("Hostname", host),
        ("IP Address", ip),
        ("Cribl Version", version),
        ("OS", os),
        ("CPU Cores", cpu),
        ("Total Memory in GBs", str(memory)),
        ("Cribl Home", file)
    )
    document_obj.add_heading("Leader Node", 2)
    document_obj.add_paragraph("System Specifications")
    leader_table=document_obj.add_table(rows=0, cols=2)
    table_by_record(leader_table, leader_node_information, "Medium Grid 1 Accent 1")
    document_obj.add_heading("Leader Node Failover", 2)
    document_obj.add_paragraph("<DESCRIPTION OF HA/DR SETUP>")

# Creates a table containing information about the Firewall Rules/Ports for the Leader Node
def leader_firewall():
    document_obj.add_heading("Firewall Rules/Ports(Leader)", 1)
    document_obj.add_heading("Leader", 2)
    # 
    firewall_table_information = [("Name", "Source", "Port")]
    for port_information in api_return["Firewall Rules/Ports(Leader)"]:
        name = port_information["Name"]
        source = port_information["Source"]
        port = port_information["Port"]
        entry = name,source,str(port)
        firewall_table_information.append(entry)
    leaderfirewall_table=document_obj.add_table(rows=0,cols=3)
    table_by_record(leaderfirewall_table, firewall_table_information, "Colorful Grid Accent 1")

# Creates all tables that can be automatically populated with the JSON information
def create_all_autofill_tables():
    # Iterates through JSON
    leader_firewall_added = False
    for entry in api_return:
        if entry not in ("Leader Node", "Workers", "Firewall Rules/Ports(Leader)", "Git", "Syslog"):
            # Adds a heading for each dictionary entry in the JSON
            # Removes underscores left from coded output
            if entry == "Firewall Rules/Ports(Worker)" and leader_firewall_added == False:
                leader_firewall()
                leader_firewall_added = True
            # Handles underscores for appearance
            document_obj.add_heading(entry.replace("_", " "), 1)
            # Iterates through groups in each entry
            quickconnect=False
            for group in api_return[entry]:
                if(api_return[entry][group]!=None):
                    # Adds a heading for each group
                    document_obj.add_heading(group + " Worker Group", 2)
                    # Creates the populated table dependent on the entry and group
                    populate_table_by_json(entry, group)
                    quickconnect=True
            if not quickconnect:
                document_obj.add_paragraph("No "+entry+" in this deployment")
    
# Creates a table for given entry and group from API call
def populate_table_by_json(entry, group):
    # Creates template table with as many columns as pieces of information
    template_table = document_obj.add_table(rows=1, cols=len(api_return[entry][group][0]))
    template_table.style="Colorful Grid Accent 1"
    title_line = template_table.rows[0].cells
    # Populates title line with column headings
    for index,key in enumerate(api_return[entry][group][0].keys()):
        title_line[index].text = key
    # Iterates through information in each entry by group
    for value_dict in api_return[entry][group]:
        # Adds new row for information
        template_table_rows = template_table.add_row().cells
        # Populates cells in correct column
        for index,value in enumerate(value_dict):
            template_table_rows[index].text = str(value_dict[value])

# Adds template Use Cases information to the document
def use_cases():
    document_obj.add_heading("Use Cases", 1)
    document_obj.add_heading("Syslog", 2)
    document_obj.add_paragraph(client_name+" will use Cribl Stream to create pipelines to reduce and/or transform data from up to ## Syslog feeds; as of today, the following Syslog sources are currently configured in the <WORKER GROUP> worker group. This use case will be deployed in the <WORKER GROUP> worker group.")
    # for group in group_names:
    #     document_obj.add_heading(group + " Worker Group", 3)
    #     populate_table_by_json("Syslog", group)

# Automates the adding of large quantities of text used as a template
def block_footer_creation():
    footer = open("Scripts/block_footer.json", "r", encoding="cp850")
    data = json.load(footer)
    for key,value in data.items():
        if key == "Best Practices & Recommendations":
            document_obj.add_heading(key, 1)
        else:
            document_obj.add_heading(key, 2)
        # Handles the addition of an image whilst retaining iteration
        if type(value) is list:
            document_obj.add_paragraph(value[0].replace("<CLIENT NAME>", client_name))
            document_obj.add_picture("Scripts/"+value[1])
        else:
            document_obj.add_paragraph(value.replace("<CLIENT NAME>", client_name))
    footer.close()

# Creates a leader load balancer table
def leader_load_balance():
    leader_load_balance_record=(("", "Leader 1", "Leader 2"),("ID", "Enter value here", "Enter value here"), ("Name", "Enter value here", "Enter value here"), ("Port", "Enter value here", "Enter value here"))
    document_obj.add_heading("Leader Load Balancing", 2)
    leader_load_balance_table = document_obj.add_table(rows=0, cols=len(leader_load_balance_record[0]))
    table_by_record(leader_load_balance_table, leader_load_balance_record, "Colorful Grid Accent 1")

if __name__ == "__main__":
    script_choice=input("Which document would you like to request: Press 1 for As Built, or 2 for Health Check(WIP): ")
    # Variable assingment
    client_name = input("Enter Client name:\n")
    if client_name == "":
        client_name = "<CLIENT NAME>"
    # Allows for custom consultant name to be input
    consultant_name = input("Enter Consultant name:\n")
    if consultant_name == "":
        consultant_name = "<CONSULTANT>"

    url=input("What is the URL for your Cribl Environment(ie. http://localhost:9000): ")
    url="http://localhost:9000" if url=="" else url

    auth_type = int(input("Would you like to log in with:\n\t1) Username and Password(Docker Instances)\n\t2) Bearer Token\n\t3) Client ID and Client Secret\n>: "))
    match auth_type:
        case 1:
            username=input("Username: ")
            username="admin" if username=="" else username
            password=input("Password: ")
            password="password" if password=="" else password

            token = requests.post(f"{url}/api/v1/auth/login", json={"username":username,"password":password})
            token_dict=token.json()["token"]
            header = {"accept": "application/json", "Authorization": f"Bearer {token_dict}"}
            # Run API request
        case 2:
            bearer = input("What is the Bearer Token (ie. eyJhbGciOi...): ")
            header={"accept": "application/json", "Authorization": f"Bearer {bearer}"}
        case 3:
            clientID=input("Client ID: ")
            clientID="admin" if clientID=="" else clientID
            clientSecret=input("Client Secret: ")
            clientSecret="password" if clientSecret=="" else clientSecret
            
            token = requests.post('https://login.cribl.cloud/oauth/token',
                                  headers={
                                        'content-type': 'application/json',
                                    },
                                  json={
                                        'grant_type': 'client_credentials',
                                        'client_id': clientID,
                                        'client_secret': clientSecret,
                                        'audience': 'https://api.cribl.cloud',
                                    })
            token_dict=token.json()["access_token"]
            header = {"accept": "application/json", "Authorization": f"Bearer {token_dict}"}

    print("Information Received!")
    print("Generating your document....")
    executive_summary_text = f"The following As Built architecture is a document meant to explain the configuration and architecture of the Cribl Stream platform within the {client_name} network, along with use case explanation and best practices. "
    # Creates main document object
    document_obj = Document()
    # Set variable to contain information from API script
    api_return = get_json.main(url,header,script_choice)
    # Collect all worker group names
    group_names = get_json.all_workgroups(header,url)
    # Populate Document
    if script_choice=="1":
        populate_document()
        # Outputs document
        document_obj.save(f"{client_name} Cribl Envrionment As-Built.docx")
    elif script_choice=="2":
        health_document()
        document_obj.save(f"{client_name} Cribl Environment Health-Check.docx")
